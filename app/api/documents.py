"""
Enhanced Document Management API with Guru Agent Integration
Supports multiple file formats, intelligent content extraction, and educational features
"""

from fastapi import APIRouter, Depends, HTTPException, status, UploadFile, File, Form
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select
from typing import List, Optional, Dict, Any
from pydantic import BaseModel
import aiofiles
import os
from pathlib import Path
import uuid
from datetime import datetime
import logging

from app.database.base import get_db
from app.database.models import Document as DocumentModel, User
from app.database.schemas import Document, DocumentCreate, DocumentUpdate
from app.auth.dependencies import get_current_active_user
from app.rag.advanced_rag_system import get_rag_system
from app.rag.advanced_rag_system import advanced_rag_system as enhanced_rag_system
from app.mcp_integration import mcp_manager, AgentType

logger = logging.getLogger(__name__)
router = APIRouter()

# Create uploads directory
UPLOAD_DIR = Path("uploads")
UPLOAD_DIR.mkdir(exist_ok=True)

# Allowed file types for Guru processing
GURU_ALLOWED_EXTENSIONS = {'.txt', '.pdf', '.docx', '.md', '.py', '.js', '.html', '.css', '.json'}
MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB

# Guru-specific models
class GuruDocumentRequest(BaseModel):
    action: str  # "explain", "generate_test", "summarize", "extract_concepts"
    specific_topic: Optional[str] = None
    difficulty_level: Optional[str] = "intermediate"

class GuruProcessingResponse(BaseModel):
    document_id: str
    action: str
    result: Dict[str, Any]
    processing_time_ms: float
    timestamp: str

class TestGenerationRequest(BaseModel):
    question_count: int = 5
    question_types: List[str] = ["multiple_choice", "short_answer"]
    difficulty_level: str = "intermediate"
    focus_topics: Optional[List[str]] = None


@router.get("/", response_model=List[Document])
async def get_documents(
    skip: int = 0,
    limit: int = 100,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_active_user)
):
    """Get user's documents"""
    result = await db.execute(
        select(DocumentModel)
        .where(DocumentModel.user_id == current_user.id)
        .offset(skip)
        .limit(limit)
    )
    documents = result.scalars().all()
    return documents


@router.post("/", response_model=Document)
async def create_document(
    document: DocumentCreate,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_active_user)
):
    """Create a new document"""
    # Create document in database
    db_document = DocumentModel(
        title=document.title,
        content=document.content,
        file_type=document.file_type,
        user_id=current_user.id,
        metadata=document.metadata
    )
    
    db.add(db_document)
    await db.commit()
    await db.refresh(db_document)
    
    # Add to vector database
    try:
        chunk_ids = await rag_system.add_document(
            document_content=document.content,
            document_id=str(db_document.id),
            metadata={
                "user_id": current_user.id,
                "title": document.title,
                "file_type": document.file_type,
                **(document.metadata or {})
            }
        )
        
        # Update document with vector IDs
        db_document.vector_id = ",".join(chunk_ids)
        await db.commit()
        
    except Exception as e:
        # If vector indexing fails, we still keep the document but log the error
        print(f"Failed to index document {db_document.id}: {e}")
    
    return db_document


@router.post("/upload", response_model=Document)
async def upload_document(
    file: UploadFile = File(...),
    title: Optional[str] = None,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_active_user)
):
    """Upload and process a document file"""
    # Validate file type
    allowed_types = {".txt", ".pdf", ".docx", ".md"}
    file_extension = Path(file.filename).suffix.lower()
    
    if file_extension not in allowed_types:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"File type {file_extension} not supported. Allowed types: {allowed_types}"
        )
    
    # Save file
    file_path = UPLOAD_DIR / f"{current_user.id}_{file.filename}"
    
    async with aiofiles.open(file_path, 'wb') as f:
        content = await file.read()
        await f.write(content)
    
    # Extract text content based on file type
    try:
        if file_extension == ".txt" or file_extension == ".md":
            async with aiofiles.open(file_path, 'r', encoding='utf-8') as f:
                file_content = await f.read()
        elif file_extension == ".pdf":
            # Import here to avoid dependency issues if not installed
            from pypdf import PdfReader
            reader = PdfReader(file_path)
            file_content = ""
            for page in reader.pages:
                file_content += page.extract_text() + "\n"
        elif file_extension == ".docx":
            from docx import Document as DocxDocument
            doc = DocxDocument(file_path)
            file_content = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        else:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Unsupported file type"
            )
    
    except Exception as e:
        # Clean up file if processing fails
        if file_path.exists():
            os.remove(file_path)
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Failed to process file: {str(e)}"
        )
    
    # Create document record
    document_title = title or file.filename
    db_document = DocumentModel(
        title=document_title,
        content=file_content,
        file_path=str(file_path),
        file_type=file_extension,
        user_id=current_user.id,
        metadata={"original_filename": file.filename}
    )
    
    db.add(db_document)
    await db.commit()
    await db.refresh(db_document)
    
    # Add to vector database
    try:
        chunk_ids = await rag_system.add_document(
            document_content=file_content,
            document_id=str(db_document.id),
            metadata={
                "user_id": current_user.id,
                "title": document_title,
                "file_type": file_extension,
                "original_filename": file.filename
            }
        )
        
        db_document.vector_id = ",".join(chunk_ids)
        await db.commit()
        
    except Exception as e:
        print(f"Failed to index uploaded document {db_document.id}: {e}")
    
    return db_document


@router.post("/upload-enhanced", response_model=Document)
async def upload_document_enhanced(
    file: UploadFile = File(...),
    title: Optional[str] = Form(None),
    category: Optional[str] = Form("general"),
    tags: Optional[str] = Form(""),  # Comma-separated tags
    auto_extract_info: bool = Form(True),
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_active_user)
):
    """
    Enhanced document upload with intelligent content extraction and categorization
    Supports: PDF, DOCX, TXT, MD files with metadata extraction
    """
    # Validate file type
    allowed_types = {".txt", ".pdf", ".docx", ".md", ".rtf"}
    file_extension = Path(file.filename).suffix.lower()
    
    if file_extension not in allowed_types:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"File type {file_extension} not supported. Allowed types: {allowed_types}"
        )
    
    # Generate unique filename to avoid conflicts
    unique_filename = f"{current_user.id}_{uuid.uuid4()}_{file.filename}"
    file_path = UPLOAD_DIR / unique_filename
    
    # Save file
    try:
        async with aiofiles.open(file_path, 'wb') as f:
            content = await file.read()
            await f.write(content)
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to save file: {str(e)}"
        )
    
    # Extract text content with enhanced processing
    try:
        file_content, extracted_metadata = await _extract_content_enhanced(file_path, file_extension)
        
        if not file_content.strip():
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="No text content could be extracted from the file"
            )
        
    except Exception as e:
        # Clean up file if processing fails
        if file_path.exists():
            os.remove(file_path)
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Failed to process file: {str(e)}"
        )
    
    # Prepare document metadata
    document_title = title or extracted_metadata.get("title", file.filename)
    
    # Parse tags
    tag_list = [tag.strip() for tag in tags.split(",") if tag.strip()] if tags else []
    
    # Enhanced metadata
    enhanced_metadata = {
        "original_filename": file.filename,
        "file_size": len(content),
        "upload_timestamp": datetime.utcnow().isoformat(),
        "category": category,
        "tags": tag_list,
        "auto_extracted": auto_extract_info,
        **extracted_metadata
    }
    
    # Create document record
    db_document = DocumentModel(
        title=document_title,
        content=file_content,
        file_path=str(file_path),
        file_type=file_extension,
        user_id=current_user.id,
        metadata=enhanced_metadata
    )
    
    db.add(db_document)
    await db.commit()
    await db.refresh(db_document)
    
    # Add to both RAG systems for comprehensive indexing
    try:
        # Add to standard RAG system
        chunk_ids = await rag_system.add_document(
            document_content=file_content,
            document_id=str(db_document.id),
            metadata={
                "user_id": current_user.id,
                "title": document_title,
                "file_type": file_extension,
                "category": category,
                **enhanced_metadata
            }
        )
        
        # Add to enhanced RAG system for user-specific knowledge
        await enhanced_rag_system.add_document(
            user_id=str(current_user.id),
            content=file_content,
            metadata={
                "document_id": str(db_document.id),
                "title": document_title,
                "category": category,
                "tags": tag_list,
                **enhanced_metadata
            }
        )
        
        # Update document with vector IDs
        db_document.vector_id = ",".join(chunk_ids)
        await db.commit()
        
        # Log this interaction for user learning
        await enhanced_rag_system.add_user_interaction(
            user_id=str(current_user.id),
            interaction={
                "action": "document_upload",
                "document_title": document_title,
                "category": category,
                "file_type": file_extension,
                "content_length": len(file_content),
                "timestamp": datetime.utcnow().isoformat()
            }
        )
        
    except Exception as e:
        print(f"Failed to index uploaded document {db_document.id}: {e}")
        # Don't fail the upload, just log the error
    
    return db_document


async def _extract_content_enhanced(file_path: Path, file_extension: str) -> tuple[str, dict]:
    """
    Enhanced content extraction with metadata extraction
    Returns: (content, metadata_dict)
    """
    content = ""
    metadata = {}
    
    try:
        if file_extension in [".txt", ".md"]:
            async with aiofiles.open(file_path, 'r', encoding='utf-8') as f:
                content = await f.read()
                
            # Extract basic metadata for text files
            lines = content.split('\n')
            metadata["line_count"] = len(lines)
            metadata["word_count"] = len(content.split())
            
            # Try to extract title from first line if it looks like a title
            if lines and (lines[0].startswith('#') or len(lines[0]) < 100):
                metadata["title"] = lines[0].strip('#').strip()
                
        elif file_extension == ".pdf":
            try:
                from pypdf import PdfReader
                reader = PdfReader(str(file_path))
                
                # Extract text
                for page in reader.pages:
                    content += page.extract_text() + "\n"
                
                # Extract metadata
                if reader.metadata:
                    metadata["pdf_title"] = reader.metadata.get('/Title', '')
                    metadata["pdf_author"] = reader.metadata.get('/Author', '')
                    metadata["pdf_subject"] = reader.metadata.get('/Subject', '')
                    metadata["pdf_creator"] = reader.metadata.get('/Creator', '')
                
                metadata["page_count"] = len(reader.pages)
                metadata["word_count"] = len(content.split())
                
            except ImportError:
                raise HTTPException(
                    status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                    detail="PDF processing not available. Please install pypdf."
                )
                
        elif file_extension == ".docx":
            try:
                from docx import Document as DocxDocument
                doc = DocxDocument(str(file_path))
                
                # Extract text
                content = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                
                # Extract metadata
                core_props = doc.core_properties
                metadata["docx_title"] = core_props.title or ""
                metadata["docx_author"] = core_props.author or ""
                metadata["docx_subject"] = core_props.subject or ""
                metadata["docx_keywords"] = core_props.keywords or ""
                
                if core_props.created:
                    metadata["created_date"] = core_props.created.isoformat()
                if core_props.modified:
                    metadata["modified_date"] = core_props.modified.isoformat()
                
                metadata["paragraph_count"] = len(doc.paragraphs)
                metadata["word_count"] = len(content.split())
                
            except ImportError:
                raise HTTPException(
                    status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                    detail="DOCX processing not available. Please install python-docx."
                )
                
        elif file_extension == ".rtf":
            # Basic RTF support (you might want to add striprtf library)
            async with aiofiles.open(file_path, 'r', encoding='utf-8') as f:
                raw_content = await f.read()
                # Simple RTF cleaning (basic approach)
                content = raw_content  # TODO: Add proper RTF parsing
                metadata["format"] = "rtf"
                
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Error extracting content: {str(e)}"
        )
    
    return content, metadata


@router.get("/{document_id}", response_model=Document)
async def get_document(
    document_id: int,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_active_user)
):
    """Get specific document"""
    result = await db.execute(
        select(DocumentModel)
        .where(
            DocumentModel.id == document_id,
            DocumentModel.user_id == current_user.id
        )
    )
    document = result.scalar_one_or_none()
    
    if not document:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Document not found"
        )
    
    return document


@router.delete("/{document_id}")
async def delete_document(
    document_id: int,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_active_user)
):
    """Delete a document"""
    result = await db.execute(
        select(DocumentModel)
        .where(
            DocumentModel.id == document_id,
            DocumentModel.user_id == current_user.id
        )
    )
    document = result.scalar_one_or_none()
    
    if not document:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Document not found"
        )
    
    # Remove from vector database
    try:
        await rag_system.delete_document(str(document.id))
    except Exception as e:
        print(f"Failed to remove document from vector database: {e}")
    
    # Remove file if exists
    if document.file_path and os.path.exists(document.file_path):
        os.remove(document.file_path)
    
    # Remove from database
    await db.delete(document)
    await db.commit()
    
    return {"message": "Document deleted successfully"}


@router.post("/search")
async def search_documents(
    query: str,
    top_k: int = 5,
    current_user: User = Depends(get_current_active_user)
):
    """Search documents using RAG"""
    try:
        # Simplified search for now
        return {
            "query": query,
            "results": [],
            "message": "Search functionality temporarily disabled"
        }
    except Exception as e:
        logger.error(f"Search error: {e}")
        raise HTTPException(status_code=500, detail="Search failed")


# ===============================================
# GURU AGENT SPECIFIC ENDPOINTS
# ===============================================

async def extract_text_from_file(file_path: str, filename: str) -> str:
    """Extract text content from uploaded file for Guru processing with PDF/DOCX support"""
    
    file_ext = os.path.splitext(filename)[1].lower()
    
    try:
        if file_ext in ['.txt', '.md']:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        
        elif file_ext == '.py':
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
                return f"Python Code:\n\n{content}"
        
        elif file_ext in ['.js', '.html', '.css', '.json']:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
                return f"Code file ({file_ext}):\n\n{content}"
        
        elif file_ext == '.pdf':
            # Extract text from PDF using PyMuPDF
            try:
                import fitz  # PyMuPDF
                doc = fitz.open(file_path)
                text_content = ""
                for page_num in range(len(doc)):
                    page = doc.load_page(page_num)
                    text_content += page.get_text()
                doc.close()
                
                if text_content.strip():
                    return f"PDF Content from {filename}:\n\n{text_content}"
                else:
                    return f"PDF file {filename} uploaded but no text content could be extracted."
            except ImportError:
                return f"PDF file {filename} uploaded but PyMuPDF library not available for text extraction."
            except Exception as pdf_error:
                logger.error(f"PDF extraction error for {filename}: {pdf_error}")
                return f"PDF file {filename} uploaded but text extraction failed: {str(pdf_error)}"
        
        elif file_ext in ['.docx', '.doc']:
            # Extract text from DOCX using python-docx
            try:
                from docx import Document
                doc = Document(file_path)
                text_content = ""
                for paragraph in doc.paragraphs:
                    text_content += paragraph.text + "\n"
                
                if text_content.strip():
                    return f"DOCX Content from {filename}:\n\n{text_content}"
                else:
                    return f"DOCX file {filename} uploaded but no text content could be extracted."
            except ImportError:
                return f"DOCX file {filename} uploaded but python-docx library not available for text extraction."
            except Exception as docx_error:
                logger.error(f"DOCX extraction error for {filename}: {docx_error}")
                return f"DOCX file {filename} uploaded but text extraction failed: {str(docx_error)}"
        
        else:
            return f"File uploaded: {filename}. Text extraction not available for {file_ext} files."
            
    except Exception as e:
        logger.error(f"Text extraction error for {filename}: {e}")
        return f"Error extracting text from {filename}: {str(e)}"

async def process_document_with_guru(document_content: str, action: str, **kwargs) -> Dict[str, Any]:
    """Process document content using Guru agent capabilities"""
    
    # Simulate Guru agent processing (in real implementation, this would call LLM)
    if action == "explain":
        topic = kwargs.get("specific_topic", "the document content")
        return {
            "explanation": f"As Guru (गुरु), I'll explain {topic} from the document.",
            "main_concepts": [
                "Key concept 1 identified from the document",
                "Key concept 2 with practical applications",
                "Advanced topics for further study"
            ],
            "learning_path": [
                "Start with basic understanding",
                "Practice with examples",
                "Apply concepts to real projects"
            ],
            "content_preview": document_content[:500] + "..."
        }
        
    elif action == "generate_test":
        difficulty = kwargs.get("difficulty_level", "intermediate")
        return {
            "test_questions": [
                {
                    "type": "multiple_choice",
                    "question": "What is the main topic of this document?",
                    "options": ["A) Programming", "B) Theory", "C) Practice", "D) All of the above"],
                    "correct_answer": "D"
                },
                {
                    "type": "short_answer", 
                    "question": "Explain the key concept mentioned in the document.",
                    "sample_answer": "Based on the document analysis..."
                }
            ],
            "difficulty": difficulty,
            "estimated_time": "15-20 minutes"
        }
        
    elif action == "summarize":
        return {
            "summary": "Document summary generated by Guru agent",
            "key_points": [
                "Main topic covered in the document",
                "Important concepts and definitions",
                "Practical applications and examples"
            ],
            "learning_objectives": [
                "Understand core concepts",
                "Apply knowledge practically",
                "Identify areas for improvement"
            ]
        }
        
    elif action == "extract_concepts":
        return {
            "concepts": [
                {"name": "Core Concept 1", "description": "Fundamental idea from document"},
                {"name": "Advanced Topic", "description": "Higher level concept requiring prerequisites"}
            ],
            "skills": ["Analysis", "Critical thinking", "Problem solving"],
            "prerequisites": ["Basic understanding of the subject"],
            "difficulty_assessment": kwargs.get("difficulty_level", "intermediate")
        }
        
    else:
        return {"error": f"Unknown action: {action}"}

@router.post("/guru/upload")
async def upload_document_for_guru(
    candidate_id: str = Form(...),
    file: UploadFile = File(...)
):
    """Upload a document specifically for Guru agent processing"""
    
    try:
        # Validate file
        if not file.filename:
            raise HTTPException(status_code=400, detail="No file provided")
        
        file_ext = os.path.splitext(file.filename)[1].lower()
        if file_ext not in GURU_ALLOWED_EXTENSIONS:
            raise HTTPException(
                status_code=400, 
                detail=f"File type {file_ext} not allowed for Guru processing. Supported: {', '.join(GURU_ALLOWED_EXTENSIONS)}"
            )
        
        # Check file size
        file_content = await file.read()
        if len(file_content) > MAX_FILE_SIZE:
            raise HTTPException(
                status_code=400, 
                detail=f"File too large. Max size: {MAX_FILE_SIZE // (1024*1024)}MB"
            )
        
        # Save file
        document_id = str(uuid.uuid4())
        file_path = UPLOAD_DIR / f"{document_id}_{file.filename}"
        
        async with aiofiles.open(file_path, "wb") as f:
            await f.write(file_content)
        
        # Extract text content
        extracted_text = await extract_text_from_file(str(file_path), file.filename)
        
        # Save to database (simplified - using direct database access)
        try:
            from app.database.base import AsyncSessionLocal
            async with AsyncSessionLocal() as db:
                db_document = DocumentModel(
                    id=document_id,
                    user_id=candidate_id,
                    filename=f"{document_id}_{file.filename}",
                    original_filename=file.filename,
                    file_type=file_ext,
                    file_size=len(file_content),
                    file_path=str(file_path),
                    processing_status="completed",
                    extracted_text=extracted_text[:10000],  # Limit to 10k chars
                    summary=f"Document uploaded for Guru agent processing"
                )
                db.add(db_document)
                await db.commit()
        except Exception as db_error:
            logger.warning(f"Database save failed: {db_error}")
        
        logger.info(f"Document uploaded for Guru processing: {document_id}")
        
        return {
            "document_id": document_id,
            "filename": file.filename,
            "file_size": len(file_content),
            "processing_status": "completed",
            "upload_timestamp": datetime.now().isoformat(),
            "guru_ready": True,
            "extracted_text": extracted_text  # Include extracted text in response
        }
        
    except Exception as e:
        logger.error(f"Guru document upload error: {e}")
        raise HTTPException(status_code=500, detail=f"Upload failed: {str(e)}")

@router.post("/guru/process/{document_id}")
async def process_document_with_guru_agent(
    document_id: str,
    request: GuruDocumentRequest
):
    """Process uploaded document with Guru agent for explanation, testing, etc."""
    
    start_time = datetime.now()
    
    try:
        # Get document (simplified file reading)
        document_files = list(UPLOAD_DIR.glob(f"{document_id}_*"))
        if not document_files:
            raise HTTPException(status_code=404, detail="Document not found")
        
        document_file = document_files[0]
        extracted_text = await extract_text_from_file(str(document_file), document_file.name)
        
        # Process with Guru agent
        processing_result = await process_document_with_guru(
            document_content=extracted_text,
            action=request.action,
            specific_topic=request.specific_topic,
            difficulty_level=request.difficulty_level
        )
        
        processing_time = (datetime.now() - start_time).total_seconds() * 1000
        
        return GuruProcessingResponse(
            document_id=document_id,
            action=request.action,
            result=processing_result,
            processing_time_ms=processing_time,
            timestamp=datetime.now().isoformat()
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Guru document processing error: {e}")
        raise HTTPException(status_code=500, detail=f"Processing failed: {str(e)}")

@router.post("/guru/generate-test/{document_id}")
async def generate_test_from_document(
    document_id: str,
    request: TestGenerationRequest
):
    """Generate educational test from uploaded document using Guru agent"""
    
    try:
        # Get document (simplified file reading)
        document_files = list(UPLOAD_DIR.glob(f"{document_id}_*"))
        if not document_files:
            raise HTTPException(status_code=404, detail="Document not found")
        
        document_file = document_files[0]
        extracted_text = await extract_text_from_file(str(document_file), document_file.name)
        
        # Generate test with Guru agent
        test_result = await process_document_with_guru(
            document_content=extracted_text,
            action="generate_test",
            difficulty_level=request.difficulty_level,
            question_count=request.question_count
        )
        
        test_id = str(uuid.uuid4())
        
        return {
            "document_id": document_id,
            "test_id": test_id,
            "questions": test_result.get("test_questions", []),
            "difficulty_level": request.difficulty_level,
            "estimated_time_minutes": request.question_count * 3,
            "focus_topics": request.focus_topics or ["General content"],
            "guru_analysis": test_result
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Test generation error: {e}")
        raise HTTPException(status_code=500, detail=f"Test generation failed: {str(e)}")

@router.get("/guru/documents/{candidate_id}")
async def get_guru_documents(candidate_id: str):
    """Get all documents uploaded for Guru processing by a candidate"""
    
    try:
        # Get documents from uploads directory
        candidate_documents = []
        for file_path in UPLOAD_DIR.glob("*"):
            if file_path.is_file():
                # Extract candidate info from filename if needed
                candidate_documents.append({
                    "document_id": file_path.stem.split('_')[0],
                    "filename": '_'.join(file_path.name.split('_')[1:]),
                    "file_size": file_path.stat().st_size,
                    "upload_time": datetime.fromtimestamp(file_path.stat().st_mtime).isoformat(),
                    "file_type": file_path.suffix
                })
        
        return {
            "candidate_id": candidate_id,
            "documents": candidate_documents,
            "guru_features": [
                "Document explanation",
                "Test generation", 
                "Content summarization",
                "Concept extraction"
            ]
        }
        
    except Exception as e:
        logger.error(f"Error getting Guru documents: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to get documents: {str(e)}")

@router.get("/guru/health")
async def guru_document_service_health():
    """Health check for Guru document processing service"""
    
    return {
        "status": "healthy",
        "service": "guru_document_processor",
        "features": [
            "Document upload and text extraction",
            "Content explanation with educational focus",
            "Automated test generation",
            "Concept extraction and analysis",
            "Learning path recommendations"
        ],
        "allowed_extensions": list(GURU_ALLOWED_EXTENSIONS),
        "max_file_size_mb": MAX_FILE_SIZE // (1024 * 1024),
        "agent": "Guru (गुरु) - Learning Mentor",
        "timestamp": datetime.now().isoformat()
    }
